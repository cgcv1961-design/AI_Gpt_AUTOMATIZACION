"""
AI_GPT_AUTOMATIZACION/utils/generador_word_audiovisual.py
---------------------------------------------------------

Generador de reporte Word para la vertical AUDIOVISUAL.

OBJETIVO
--------
Alinear la salida audiovisual con un formato más ejecutivo, limpio
y menos redundante, manteniendo:

- claridad de lectura
- identificación clara
- partes visibles humanas
- separación entre:
    1. severidad del contrato
    2. riesgo para la parte analizada
    3. riesgo para la contraparte

PRINCIPIO
---------
El Word se genera EXCLUSIVAMENTE desde el JSON final.

MEJORAS DE ESTA VERSIÓN
-----------------------
1. Se elimina el bloque "Cómo leer este informe".
2. Se evita duplicar duración y detalle del plazo cuando son equivalentes.
3. Se elimina el bloque separado "Qué significa este resultado", porque repetía
   demasiado el nivel ya visible en severidad y riesgo.
4. Se mantiene el cuerpo ejecutivo corto:
   - Identificación
   - Información General
   - Evaluación General
   - Resumen Ejecutivo
   - Puntos Críticos
   - Recomendación
   - Preguntas
   - Conclusión
   - Scoring
   - Anexo Técnico
5. El detalle largo queda concentrado en el Anexo Técnico.
"""

import os
from typing import Any, Dict, List
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING


def generar_word_audiovisual(resultado: dict, ruta_json: str) -> str:
    doc = Document()

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(12)

    def compactar_parrafo(parrafo, indent=0):
        pf = parrafo.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if indent:
            pf.left_indent = Pt(indent)

    def heading_compacto(texto, level=1):
        p = doc.add_heading(texto, level=level)
        compactar_parrafo(p)
        return p

    def parrafo_compacto(texto="", bold=False, indent=0):
        p = doc.add_paragraph()
        compactar_parrafo(p, indent=indent)
        run = p.add_run("" if texto is None else str(texto))
        run.bold = bold
        return p

    def parrafo_bulleted(texto="", indent=0):
        p = doc.add_paragraph(style="List Bullet")
        compactar_parrafo(p, indent=indent)
        p.add_run("" if texto is None else str(texto))
        return p

    def valor_no_vacio(*valores, default="-"):
        for v in valores:
            if v not in (None, "", [], {}):
                return v
        return default

    def texto_limpio(valor, default="-") -> str:
        if valor in (None, "", [], {}):
            return default
        return str(valor).strip()

    def lista_desde_valor(valor) -> List[Any]:
        if valor in (None, "", [], {}):
            return []
        if isinstance(valor, list):
            return [x for x in valor if x not in (None, "", [], {})]
        if isinstance(valor, dict):
            return [valor]
        return [valor]

    def formatear_visible(valor: Any) -> str:
        if isinstance(valor, dict):
            nombre = texto_limpio(valor.get("nombre") or valor.get("parte") or valor.get("name"), default="-")
            rol = texto_limpio(valor.get("rol") or valor.get("tipo") or valor.get("role"), default="")
            if nombre != "-" and rol:
                return f"{nombre} ({rol})"
            if nombre != "-":
                return nombre
            return texto_limpio(rol, default="-")
        return texto_limpio(valor, default="-")

    def formatear_duracion_audiovisual(nucleo: Dict[str, Any]) -> Dict[str, str]:
        duracion_texto = texto_limpio(
            valor_no_vacio(
                nucleo.get("duracion_texto"),
                nucleo.get("plazo_texto"),
                nucleo.get("duracion"),
                default="-"
            ),
            default="-"
        )

        if duracion_texto != "-":
            breve = duracion_texto
        else:
            duracion_meses = valor_no_vacio(nucleo.get("duracion_meses"), default=None)
            if duracion_meses not in (None, "", [], {}):
                breve = f"{duracion_meses} meses"
            else:
                duracion_dias = valor_no_vacio(nucleo.get("duracion_dias"), default=None)
                unidad_duracion = texto_limpio(nucleo.get("unidad_duracion"), default="días")
                if duracion_dias not in (None, "", [], {}):
                    breve = f"{duracion_dias} {unidad_duracion}"
                else:
                    breve = "-"

        detalle = duracion_texto if duracion_texto != "-" else ""

        return {
            "breve": breve,
            "detalle": detalle,
        }

    def formatear_precio(nucleo: Dict[str, Any]) -> str:
        precio_total = texto_limpio(
            valor_no_vacio(
                nucleo.get("precio_total"),
                nucleo.get("precio"),
                nucleo.get("monto"),
                default="-"
            ),
            default="-"
        )

        moneda = texto_limpio(
            valor_no_vacio(
                nucleo.get("moneda"),
                nucleo.get("divisa"),
                default="-"
            ),
            default="-"
        )

        if precio_total == "-" and moneda == "-":
            return "-"
        if precio_total == "-":
            return moneda
        if moneda == "-":
            return precio_total

        return f"{precio_total} {moneda}"

    def construir_distribucion_severidad(metricas: Dict[str, Any]) -> str:
        bloques = [
            ("críticos", metricas.get("riesgos_criticos", 0)),
            ("altos", metricas.get("riesgos_altos", 0)),
            ("medio-altos", metricas.get("riesgos_media_altos", 0)),
            ("medios", metricas.get("riesgos_medios", 0)),
            ("bajos", metricas.get("riesgos_bajos", 0)),
        ]

        partes = []
        for nombre, valor in bloques:
            try:
                n = int(float(valor))
            except (TypeError, ValueError):
                continue
            if n > 0:
                partes.append(f"{nombre} {n}")

        return ", ".join(partes) if partes else "sin observaciones clasificadas"

    def textos_son_equivalentes(a: str, b: str) -> bool:
        a_norm = texto_limpio(a, default="").lower()
        b_norm = texto_limpio(b, default="").lower()
        return a_norm == b_norm and a_norm != ""

    nucleo = resultado.get("nucleo_contractual", {}) or {}
    analisis_sectorial = resultado.get("analisis_sectorial", {}) or {}
    informe_cliente = resultado.get("informe_cliente", {}) or {}
    metadata = resultado.get("metadata_sistema", {}) or {}
    scoring = resultado.get("scoring", {}) or {}

    resumen = informe_cliente.get("resumen_ejecutivo", {}) or {}
    detalle = informe_cliente.get("informe_detallado", {}) or {}
    confianza = analisis_sectorial.get("nivel_confianza_analisis", {}) or {}
    metadata_presentacion = metadata.get("metadata_presentacion", {}) or {}

    tipo_contrato = texto_limpio(
        valor_no_vacio(
            nucleo.get("tipo_contrato"),
            nucleo.get("clase_contrato"),
            nucleo.get("naturaleza_contrato"),
            default="-"
        ),
        default="-"
    )

    partes = [formatear_visible(x) for x in lista_desde_valor(nucleo.get("partes"))]
    partes_con_rol = [formatear_visible(x) for x in lista_desde_valor(metadata_presentacion.get("partes_con_rol"))]
    if not partes_con_rol:
        partes_con_rol = partes

    duracion = formatear_duracion_audiovisual(nucleo)
    precio = formatear_precio(nucleo)

    riesgos = analisis_sectorial.get("riesgos_sectoriales", [])
    if not isinstance(riesgos, list):
        riesgos = []

    score_total = texto_limpio(scoring.get("score_total"), default="-")
    nivel_riesgo = texto_limpio(scoring.get("nivel_riesgo"), default="-")
    version_scoring = texto_limpio(scoring.get("version_scoring"), default="-")
    metricas = scoring.get("metricas", {}) or {}

    severidad_contrato = scoring.get("severidad_contrato", {}) or {}
    riesgo_parte = scoring.get("riesgo_parte_analizada", {}) or {}
    riesgo_contraparte = scoring.get("riesgo_contraparte", {}) or {}

    severidad_contrato_score = texto_limpio(severidad_contrato.get("score"), default=score_total)
    severidad_contrato_nivel = texto_limpio(severidad_contrato.get("nivel"), default=nivel_riesgo)

    riesgo_parte_score = texto_limpio(riesgo_parte.get("score"), default=score_total)
    riesgo_parte_nivel = texto_limpio(riesgo_parte.get("nivel"), default=nivel_riesgo)
    riesgo_parte_rol = texto_limpio(riesgo_parte.get("rol"), default="Parte analizada")

    riesgo_contraparte_score = texto_limpio(riesgo_contraparte.get("score"), default=score_total)
    riesgo_contraparte_nivel = texto_limpio(riesgo_contraparte.get("nivel"), default=nivel_riesgo)
    riesgo_contraparte_rol = texto_limpio(riesgo_contraparte.get("rol"), default="Contraparte")

    perspectiva = texto_limpio(metadata.get("perspectiva_analisis"), default="proveedor").lower()
    pais_referencia = texto_limpio(metadata.get("pais_referencia"), default="internacional").lower()

    parte_analizada_label = texto_limpio(
        metadata_presentacion.get("parte_analizada_label"),
        default="Parte analizada"
    )
    rol_contractual_detectado = texto_limpio(
        metadata_presentacion.get("rol_contractual_detectado"),
        default="Parte analizada"
    )
    nombre_parte_analizada = formatear_visible(
        metadata_presentacion.get("nombre_parte_analizada")
    )

    mapa_paises = {
        "argentina": "Argentina",
        "uruguay": "Uruguay",
        "italia": "Italia",
        "espana": "España",
        "internacional": "Internacional / Otro",
    }
    texto_pais = mapa_paises.get(pais_referencia, "Internacional / Otro")

    vision_general = texto_limpio(resumen.get("vision_general"), default="-")
    nivel_riesgo_global = texto_limpio(resumen.get("nivel_riesgo_global"), default="-")
    puntos_criticos = [
        texto_limpio(x, default="")
        for x in lista_desde_valor(resumen.get("puntos_criticos"))
        if texto_limpio(x, default="")
    ]
    preguntas = [
        texto_limpio(x, default="")
        for x in lista_desde_valor(detalle.get("preguntas_clave_antes_de_firmar"))
        if texto_limpio(x, default="")
    ]
    conclusion = texto_limpio(detalle.get("conclusion_profesional"), default="-")
    recomendacion = texto_limpio(informe_cliente.get("recomendacion_profesional"), default="-")

    heading_compacto("Reporte de Análisis Contractual Audiovisual", level=0)

    # =====================================================
    # IDENTIFICACIÓN
    # =====================================================
    heading_compacto("Identificación clara", level=1)
    parrafo_compacto(f"Usted está analizando este contrato como: {parte_analizada_label}")
    parrafo_compacto(f"Rol contractual detectado: {rol_contractual_detectado}")
    parrafo_compacto(f"Parte analizada identificada: {nombre_parte_analizada}")

    # =====================================================
    # INFORMACIÓN GENERAL
    # =====================================================
    heading_compacto("Información General", level=1)
    parrafo_compacto(f"Tipo de contrato: {tipo_contrato}")
    parrafo_compacto(f"Duración: {duracion['breve']}")

    if duracion["detalle"] and not textos_son_equivalentes(duracion["breve"], duracion["detalle"]):
        parrafo_compacto("Detalle del plazo:", bold=True)
        parrafo_compacto(duracion["detalle"], indent=18)

    parrafo_compacto(f"Precio total: {precio}")
    parrafo_compacto(f"País / contexto legal de referencia: {texto_pais}")

    # =====================================================
    # PARTES
    # =====================================================
    heading_compacto("Partes del contrato", level=1)
    if partes_con_rol:
        for idx, parte in enumerate(partes_con_rol, start=1):
            parrafo_compacto(f"Parte {idx}: {parte}")
    else:
        parrafo_compacto("No se pudieron identificar claramente las partes.")

    # =====================================================
    # EVALUACIÓN GENERAL
    # =====================================================
    heading_compacto("Evaluación General del Contrato", level=1)
    parrafo_compacto(f"Severidad del contrato: {severidad_contrato_score} ({severidad_contrato_nivel})")
    parrafo_compacto(f"Riesgo para la parte analizada ({riesgo_parte_rol}): {riesgo_parte_score} ({riesgo_parte_nivel})")
    parrafo_compacto(f"Riesgo para la contraparte ({riesgo_contraparte_rol}): {riesgo_contraparte_score} ({riesgo_contraparte_nivel})")
    parrafo_compacto(f"Cantidad de observaciones: {metricas.get('cantidad_riesgos', '-')}")
    parrafo_compacto(f"Distribución por severidad: {construir_distribucion_severidad(metricas)}")

    if nivel_riesgo_global != "-":
        parrafo_compacto(f"Nivel de riesgo global informado: {nivel_riesgo_global}")

    # =====================================================
    # RESUMEN EJECUTIVO
    # =====================================================
    heading_compacto("Resumen Ejecutivo para la Parte Analizada", level=1)
    if vision_general != "-":
        parrafo_compacto(vision_general)
    else:
        parrafo_compacto("No se generó resumen ejecutivo.")

    # =====================================================
    # PUNTOS CRÍTICOS
    # =====================================================
    heading_compacto("Puntos Críticos Principales", level=1)
    if puntos_criticos:
        for p in puntos_criticos[:3]:
            parrafo_bulleted(p)
    else:
        parrafo_compacto("No se reportaron puntos críticos principales.")

    # =====================================================
    # RECOMENDACIÓN
    # =====================================================
    heading_compacto("Recomendación Profesional", level=1)
    if recomendacion != "-":
        parrafo_compacto(recomendacion)
    else:
        parrafo_compacto("No se generó recomendación profesional.")

    # =====================================================
    # PREGUNTAS
    # =====================================================
    heading_compacto("Preguntas Clave Antes de Firmar", level=1)
    if preguntas:
        for pregunta in preguntas[:5]:
            parrafo_bulleted(pregunta)
    else:
        parrafo_compacto("No se reportaron preguntas clave antes de firmar.")

    # =====================================================
    # CONCLUSIÓN
    # =====================================================
    heading_compacto("Conclusión Profesional", level=1)
    if conclusion != "-":
        parrafo_compacto(conclusion)
    else:
        parrafo_compacto("No se reportó conclusión profesional.")

    # =====================================================
    # CONFIANZA
    # =====================================================
    heading_compacto("Nivel de Confianza del Análisis", level=1)
    parrafo_compacto(f"Nivel general: {texto_limpio(confianza.get('general'), default='-')}")
    fundamento_confianza = texto_limpio(confianza.get("fundamento"), default="-")
    if fundamento_confianza != "-":
        parrafo_compacto("Fundamento:", bold=True)
        parrafo_compacto(fundamento_confianza, indent=18)

    # =====================================================
    # SISTEMA
    # =====================================================
    heading_compacto("Sistema de Análisis Utilizado", level=1)
    parrafo_compacto("Motor de Inteligencia Artificial", bold=True)
    parrafo_compacto(f"Modelo utilizado: {texto_limpio(metadata.get('modelo_utilizado'), default='-')}")
    parrafo_compacto(f"Perfil de análisis: {texto_limpio(metadata.get('perfil_canonico'), default='-')}")
    parrafo_compacto(f"Perspectiva configurada internamente: {perspectiva}")
    parrafo_compacto(f"País de referencia: {texto_pais}")

    parrafo_compacto("Motor Jurídico-Contractual", bold=True)
    parrafo_compacto(f"Versión del motor: {texto_limpio(metadata.get('version_servicio'), default='-')}")
    parrafo_compacto(f"Versión del scoring: {version_scoring}")

    # =====================================================
    # SCORING
    # =====================================================
    heading_compacto("Resultados del Scoring", level=1)

    tabla_scoring = doc.add_table(rows=1, cols=2)
    tabla_scoring.style = "Table Grid"

    encabezado = tabla_scoring.rows[0].cells
    encabezado[0].text = "Indicador"
    encabezado[1].text = "Valor"

    filas = [
        ("Severidad del contrato", f"{severidad_contrato_score} ({severidad_contrato_nivel})"),
        (f"Riesgo para la parte analizada ({riesgo_parte_rol})", f"{riesgo_parte_score} ({riesgo_parte_nivel})"),
        (f"Riesgo para la contraparte ({riesgo_contraparte_rol})", f"{riesgo_contraparte_score} ({riesgo_contraparte_nivel})"),
        ("Cantidad de riesgos", str(metricas.get("cantidad_riesgos", "-"))),
    ]

    metricas_dinamicas = [
        ("Riesgos críticos", metricas.get("riesgos_criticos", 0)),
        ("Riesgos altos", metricas.get("riesgos_altos", 0)),
        ("Riesgos medio-altos", metricas.get("riesgos_media_altos", 0)),
        ("Riesgos medios", metricas.get("riesgos_medios", 0)),
        ("Riesgos bajos", metricas.get("riesgos_bajos", 0)),
    ]

    for nombre, valor in metricas_dinamicas:
        try:
            n = int(float(valor))
        except (TypeError, ValueError):
            continue
        if n > 0:
            filas.append((nombre, str(n)))

    for indicador, valor in filas:
        row = tabla_scoring.add_row().cells
        row[0].text = indicador
        row[1].text = valor

    # =====================================================
    # ANEXO TÉCNICO
    # =====================================================
    heading_compacto("Anexo Técnico - Detalle Ampliado de Riesgos", level=1)

    if not riesgos:
        parrafo_compacto("No se detectaron riesgos sectoriales.")
    else:
        for i, r in enumerate(riesgos, start=1):
            severidad = texto_limpio(r.get("severidad"), default="-").capitalize()
            impacto = texto_limpio(r.get("impacto"), default="-")
            descripcion = texto_limpio(r.get("descripcion"), default="-")
            recomendacion_riesgo = texto_limpio(r.get("recomendacion"), default="-")

            parrafo_compacto(f"Riesgo {i}", bold=True)

            if severidad != "-":
                parrafo_compacto(f"Severidad: {severidad}")

            if impacto != "-":
                parrafo_compacto(f"Impacto: {impacto}")

            parrafo_compacto("Descripción:", bold=True)
            parrafo_compacto(descripcion, indent=18)

            if recomendacion_riesgo != "-":
                parrafo_compacto("Sugerencia / mitigación:", bold=True)
                parrafo_compacto(recomendacion_riesgo, indent=18)

    parrafo_compacto(
        "Este informe fue generado mediante un sistema automatizado de análisis contractual "
        "basado en inteligencia artificial y un motor de evaluación jurídica propietario."
    )

    os.makedirs("output", exist_ok=True)
    nombre_docx = os.path.basename(ruta_json).replace(".json", ".docx")
    ruta_docx = os.path.join("output", nombre_docx)
    doc.save(ruta_docx)

    return ruta_docx