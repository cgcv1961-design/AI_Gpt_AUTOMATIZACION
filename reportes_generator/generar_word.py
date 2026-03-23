
"""
AI_GPT_AUTOMATIZACION/reportes_generator/generar_word.py

Genera archivo word a partir de salida json de FastAPI

Este script:
    Lee JSON
    Construye narrativa clara
    Ordena riesgos por severidad
    Genera Word profesional
    Guarda archivo listo

"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import json
import os


def generar_informe_word(ruta_json, nombre_salida="informe_analisis_contractual.docx"):
    """
    Genera un informe Word profesional a partir del JSON
    producido por el sistema de análisis contractual.
    """

    # Cargar JSON
    with open(ruta_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    document = Document()

    # =========================
    # PORTADA
    # =========================
    document.add_heading("INFORME DE ANÁLISIS CONTRACTUAL", level=1)

    fecha_actual = datetime.now().strftime("%d/%m/%Y")

    tipo_contrato = data.get("nucleo_contractual", {}).get("tipo_contrato", "No especificado")
    vertical = data.get("metadata_sistema", {}).get("vertical", "No especificada")
    perfil = data.get("metadata_sistema", {}).get("perfil_canonico", "No especificado")

    nivel = data.get("evaluacion_general", {}).get("score_riesgo", {}).get("nivel", "")
    score = data.get("evaluacion_general", {}).get("score_riesgo", {}).get("valor", "")
    version_motor = data.get("evaluacion_general", {}).get("version_scoring", "")

    document.add_paragraph(f"Tipo de contrato: {tipo_contrato}")
    document.add_paragraph(f"Vertical aplicada: {vertical}")
    document.add_paragraph(f"Perfil evaluado: {perfil}")
    document.add_paragraph(f"Fecha de análisis: {fecha_actual}")
    document.add_paragraph("")
    document.add_paragraph(f"Nivel de riesgo: {nivel.upper()}")
    document.add_paragraph(f"Score total: {score}")

    document.add_page_break()

    # =========================
    # RESUMEN EJECUTIVO
    # =========================
    document.add_heading("1. Resumen Ejecutivo", level=2)

    resumen = data.get("informe_cliente", {}).get("resumen_ejecutivo", "")
    recomendacion = data.get("informe_cliente", {}).get("recomendacion_profesional", "")

    document.add_paragraph(resumen)
    document.add_paragraph("")
    document.add_paragraph("Recomendación profesional:")
    document.add_paragraph(recomendacion)

    # =========================
    # RESULTADO DEL ANÁLISIS
    # =========================
    document.add_heading("2. Resultado del Análisis", level=2)

    document.add_paragraph(f"Nivel de riesgo determinado: {nivel.upper()}")
    document.add_paragraph(f"Score total obtenido: {score}")

    document.add_paragraph(
        "El resultado surge de la aplicación de un sistema híbrido que combina "
        "estructuración mediante inteligencia artificial y un motor determinístico "
        "de evaluación sectorial."
    )

    # =========================
    # MAPA DE RIESGOS
    # =========================
    document.add_heading("3. Mapa de Riesgos Detectados", level=2)

    riesgos = data.get("analisis_sectorial", {}).get("riesgos_sectoriales", [])

    # Ordenar por severidad
    orden_severidad = {"critica": 4, "alta": 3, "media": 2, "baja": 1}
    riesgos_ordenados = sorted(
        riesgos,
        key=lambda r: orden_severidad.get(r.get("severidad", "baja"), 0),
        reverse=True
    )

    table = document.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "N°"
    hdr[1].text = "Descripción"
    hdr[2].text = "Impacto"
    hdr[3].text = "Severidad"

    for i, riesgo in enumerate(riesgos_ordenados, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = riesgo.get("descripcion", "")
        row[2].text = riesgo.get("impacto", "")
        row[3].text = riesgo.get("severidad", "")

    # =========================
    # FUNDAMENTO DEL NIVEL
    # =========================
    document.add_heading("4. Fundamento del Nivel de Riesgo", level=2)

    if nivel == "bajo":
        texto_nivel = (
            "El contrato presenta un nivel de riesgo limitado, con observaciones menores "
            "que no comprometen de forma sustancial su viabilidad."
        )
    elif nivel == "medio":
        texto_nivel = (
            "El contrato presenta debilidades estructurales relevantes que requieren "
            "revisión y ajuste antes de su formalización, aunque no evidencian "
            "condiciones críticas irreversibles."
        )
    elif nivel == "alto":
        texto_nivel = (
            "El contrato contiene condiciones que pueden generar riesgos significativos "
            "desde el punto de vista legal o económico, por lo que se recomienda "
            "una revisión sustancial antes de su ejecución."
        )
    else:
        texto_nivel = (
            "El contrato presenta riesgos estructurales de alta criticidad que pueden "
            "comprometer seriamente su validez o viabilidad."
        )

    document.add_paragraph(texto_nivel)

    # =========================
    # METODOLOGÍA
    # =========================
    document.add_heading("5. Metodología Aplicada", level=2)

    document.add_paragraph(
        "Este informe fue generado mediante un sistema híbrido compuesto por:"
    )
    document.add_paragraph("- Modelo de lenguaje para estructuración contractual.")
    document.add_paragraph("- Validación automatizada de esquema.")
    document.add_paragraph("- Motor determinístico de scoring sectorial.")
    document.add_paragraph(f"- Versión del motor aplicado: {version_motor}")

    document.add_paragraph(
        "El presente informe constituye una herramienta de apoyo al análisis "
        "y no sustituye el asesoramiento profesional personalizado."
    )

    # Guardar archivo
    document.save(nombre_salida)

    print(f"Informe generado correctamente: {nombre_salida}")